import os
from pathlib import Path
from typing import List, Dict, Any, Optional, Set
import docx
import hashlib
from datetime import datetime
import logging
import json

import chromadb
from chromadb.config import Settings

from config import Config
from utils.local_model import LlamaModel
from utils.embeddings import CustomEmbeddingFunction

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self,
                 docs_dir: str = None,
                 chroma_dir: str = None,
                 model: Optional[LlamaModel] = None,
                 skip_processing: bool = False):
        """
        Инициализация процессора документов с возможностью пропуска обработки
        """
        self.docs_dir = Path(docs_dir or Config.DOCS_DIR)
        self.chroma_dir = Path(chroma_dir or Config.CHROMA_DIR)
        self.skip_processing = skip_processing

        # Создаем папки если они не существуют
        self.docs_dir.mkdir(parents=True, exist_ok=True)
        self.chroma_dir.mkdir(parents=True, exist_ok=True)

        if model is not None:
            self.model = model
            logger.info(f"Используется переданная модель: {self.model.embedding_model}")
        else:
            self.model = LlamaModel(
                model_name="llama3.2:3b",
                embedding_model="nomic-embed-text"
            )

        # Инициализируем кастомную функцию эмбеддингов
        self.embedding_function = CustomEmbeddingFunction(self.model)

        # Инициализируем ChromaDB
        self.chroma_client = chromadb.PersistentClient(
            path=str(self.chroma_dir),
            settings=Settings(
                anonymized_telemetry=False,
                is_persistent=True
            )
        )

        # Создаем или получаем коллекцию
        self.collection_name = "documents"
        self.collection = self._get_or_create_collection()

        # Кэш для хешей обработанных файлов
        self.processed_files_cache = self.chroma_dir / "processed_files.json"
        self._load_processed_cache()

        logger.info(f"DocumentProcessor инициализирован")

    def _load_processed_cache(self):
        """Загрузка кэша обработанных файлов"""
        self.processed_files = {}
        if self.processed_files_cache.exists():
            try:
                with open(self.processed_files_cache, 'r') as f:
                    self.processed_files = json.load(f)
                logger.info(f"Загружен кэш {len(self.processed_files)} обработанных файлов")
            except Exception as e:
                logger.warning(f"Не удалось загрузить кэш файлов: {e}")

    def _save_processed_cache(self):
        """Сохранение кэша обработанных файлов"""
        try:
            with open(self.processed_files_cache, 'w') as f:
                json.dump(self.processed_files, f, indent=2)
        except Exception as e:
            logger.error(f"Не удалось сохранить кэш файлов: {e}")

    def _get_or_create_collection(self):
        """Создает или получает коллекцию с правильной функцией эмбеддингов"""
        try:
            # Пытаемся получить существующую коллекцию
            collection = self.chroma_client.get_collection(self.collection_name)
            logger.info(f"Найдена существующая коллекция: {self.collection_name}")
            return collection
        except Exception as e:
            logger.info(f"Создаем новую коллекцию: {self.collection_name}")

            collection = self.chroma_client.create_collection(
                name=self.collection_name,
                embedding_function=self.embedding_function,
                metadata={
                    "description": "Хранилище векторных представлений документов",
                    "embedding_model": self.model.embedding_model,
                    "chunk_size": Config.CHUNK_SIZE,
                    "created_at": datetime.now().isoformat()
                }
            )
            return collection

    def process_all_documents_incremental(self) -> List[Dict[str, Any]]:
        """
        Инкрементальная обработка только новых или измененных документов
        """
        if self.skip_processing:
            logger.info("Пропускаем обработку документов (skip_processing=True)")
            return []

        processed_docs = []
        docx_files = list(self.docs_dir.glob("*.docx"))

        if not docx_files:
            logger.warning(f"В папке {self.docs_dir} не найдено файлов .docx")
            return processed_docs

        logger.info(f"Найдено {len(docx_files)} файлов .docx. Проверяем изменения...")

        for file_path in docx_files:
            if file_path.name.startswith('~$'):
                continue

            # Проверяем, нужно ли обрабатывать файл
            if self._should_process_file(file_path):
                doc_info = self.process_single_document(file_path)
                if doc_info:
                    processed_docs.append(doc_info)
                    # Обновляем кэш
                    self.processed_files[file_path.name] = {
                        'hash': doc_info["file_hash"],
                        'size': doc_info["file_size"],
                        'modified': doc_info["last_modified"].isoformat(),
                        'processed_at': datetime.now().isoformat()
                    }
            else:
                logger.debug(f"Файл {file_path.name} уже обработан, пропускаем")

        if processed_docs:
            self._save_processed_cache()
            logger.info(f"Обработано {len(processed_docs)} новых/измененных документов")
        else:
            logger.info("Нет новых или измененных документов для обработки")

        return processed_docs

    def _should_process_file(self, file_path: Path) -> bool:
        """
        Проверяет, нужно ли обрабатывать файл
        """
        file_stat = file_path.stat()
        file_hash = self.get_file_hash(file_path)

        # Если файл уже в кэше
        if file_path.name in self.processed_files:
            cache_info = self.processed_files[file_path.name]

            # Проверяем хеш, размер и дату изменения
            if (cache_info.get('hash') == file_hash and
                    cache_info.get('size') == file_stat.st_size and
                    cache_info.get('modified') == datetime.fromtimestamp(file_stat.st_mtime).isoformat()):
                return False

        return True
    # Остальные методы остаются без изменений...
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Извлечение текста из docx файла

        Args:
            file_path: Путь к файлу

        Returns:
            Текст документа
        """
        try:
            doc = docx.Document(file_path)
            full_text = []

            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            # Извлекаем текст из заголовков
            for element in doc.element.body:
                if element.tag.endswith('title'):
                    for paragraph in element.iter():
                        if paragraph.text:
                            full_text.append(paragraph.text.strip())

            return "\n".join(full_text)
        except Exception as e:
            logger.error(f"Ошибка при чтении файла {file_path}: {e}")
            return ""

    def split_text(self,
                   text: str,
                   chunk_size: int = None,
                   overlap: int = None) -> List[str]:
        """
        Умное разделение текста на чанки с учетом структуры

        Args:
            text: Исходный текст
            chunk_size: Размер чанка в символах
            overlap: Перекрытие между чанками

        Returns:
            Список текстовых чанков
        """
        chunk_size = chunk_size or Config.CHUNK_SIZE
        overlap = overlap or Config.CHUNK_OVERLAP

        if not text or len(text.strip()) == 0:
            return []

        if len(text) <= chunk_size:
            return [text.strip()]

        chunks = []
        paragraphs = text.split('\n')
        current_chunk = []
        current_length = 0

        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue

            paragraph_length = len(paragraph)

            # Если параграф слишком большой, разбиваем его
            if paragraph_length > chunk_size:
                # Если есть накопленный чанк, сохраняем его
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))
                    current_chunk = []
                    current_length = 0

                # Разбиваем большой параграф на части
                words = paragraph.split()
                temp_chunk = []
                temp_length = 0

                for word in words:
                    word_length = len(word) + 1  # +1 для пробела

                    if temp_length + word_length > chunk_size:
                        if temp_chunk:
                            chunks.append(' '.join(temp_chunk))
                            # Сохраняем перекрытие
                            temp_chunk = temp_chunk[-overlap // 10:]  # Примерно 10% перекрытия
                            temp_length = sum(len(w) + 1 for w in temp_chunk)

                    temp_chunk.append(word)
                    temp_length += word_length

                if temp_chunk:
                    chunks.append(' '.join(temp_chunk))

            # Если параграф помещается в текущий чанк
            elif current_length + paragraph_length <= chunk_size:
                current_chunk.append(paragraph)
                current_length += paragraph_length + 1  # +1 для символа новой строки

            # Если не помещается, начинаем новый чанк
            else:
                if current_chunk:
                    chunks.append('\n'.join(current_chunk))

                # Сохраняем перекрытие: берем последние несколько параграфов
                overlap_paragraphs = min(len(current_chunk), overlap // 100)  # Примерно 1% перекрытия
                current_chunk = current_chunk[-overlap_paragraphs:] if overlap_paragraphs > 0 else []
                current_chunk.append(paragraph)
                current_length = sum(len(p) + 1 for p in current_chunk)

        # Добавляем последний чанк
        if current_chunk:
            chunks.append('\n'.join(current_chunk))

        return chunks

    def get_file_hash(self, file_path: Path) -> str:
        """
        Вычисление хеша файла для отслеживания изменений

        Args:
            file_path: Путь к файлу

        Returns:
            MD5 хеш файла
        """
        try:
            hasher = hashlib.md5()
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.error(f"Ошибка при вычислении хеша файла {file_path}: {e}")
            return ""

    def process_single_document(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """
        Обработка одного документа с улучшенной информацией

        Args:
            file_path: Путь к файлу

        Returns:
            Словарь с информацией о документе
        """
        try:
            logger.info(f"Обработка документа: {file_path.name}")

            # Извлекаем текст
            text = self.extract_text_from_docx(file_path)

            if not text or len(text.strip()) < 10:  # Минимум 10 символов
                logger.warning(f"Документ {file_path.name} пуст или слишком мал")
                return None

            # Разделяем на чанки
            chunks = self.split_text(text)

            if not chunks:
                logger.warning(f"Не удалось разделить документ {file_path.name} на чанки")
                return None

            # Вычисляем хеш файла
            file_hash = self.get_file_hash(file_path)

            # Получаем статистику
            word_count = len(text.split())
            paragraph_count = len([p for p in text.split('\n') if p.strip()])

            # Информация о документе
            doc_info = {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_hash": file_hash,
                "file_size": file_path.stat().st_size,
                "last_modified": datetime.fromtimestamp(file_path.stat().st_mtime),
                "text": text,
                "chunks": chunks,
                "chunk_count": len(chunks),
                "total_chars": len(text),
                "word_count": word_count,
                "paragraph_count": paragraph_count,
                "avg_chunk_size": sum(len(c) for c in chunks) / len(chunks) if chunks else 0,
                "processed_at": datetime.now()
            }

            logger.info(f"  Извлечено: {len(chunks)} чанков, {word_count} слов, {len(text)} символов")
            return doc_info

        except Exception as e:
            logger.error(f"Ошибка при обработке документа {file_path}: {e}")
            return None

    def process_all_documents(self) -> List[Dict[str, Any]]:
        """Алиас для обратной совместимости"""
        return self.process_all_documents_incremental()

    def add_documents_to_vector_db(self, docs_info: List[Dict[str, Any]]):
        """
        Добавление документов в векторную базу с проверкой дубликатов

        Args:
            docs_info: Список информации о документах
        """
        try:
            if not docs_info:
                logger.warning("Нет документов для добавления в векторную базу")
                return

            all_ids = []
            all_texts = []
            all_metadatas = []
            skipped_count = 0

            # Проверяем существующие документы
            existing_docs = self.collection.get()
            existing_ids = set(existing_docs.get('ids', [])) if existing_docs else set()

            # Подготавливаем данные для добавления
            for doc_idx, doc_info in enumerate(docs_info):
                filename = doc_info["filename"]
                file_hash = doc_info["file_hash"]

                for chunk_idx, chunk_text in enumerate(doc_info["chunks"]):
                    # Создаем уникальный ID
                    chunk_id = f"{filename}_{file_hash}_{chunk_idx}"

                    # Пропускаем если уже существует
                    if chunk_id in existing_ids:
                        skipped_count += 1
                        continue

                    # Метаданные для чанка
                    metadata = {
                        "filename": filename,
                        "file_path": doc_info["file_path"],
                        "file_hash": file_hash,
                        "chunk_index": chunk_idx,
                        "total_chunks": doc_info["chunk_count"],
                        "source": "docx",
                        "word_count": len(chunk_text.split()),
                        "char_count": len(chunk_text),
                        "processed_at": doc_info["processed_at"].isoformat(),
                        "last_modified": doc_info["last_modified"].isoformat()
                    }

                    all_ids.append(chunk_id)
                    all_texts.append(chunk_text)
                    all_metadatas.append(metadata)

            # Добавляем в коллекцию
            if all_ids:
                self.collection.add(
                    ids=all_ids,
                    documents=all_texts,
                    metadatas=all_metadatas
                )
                logger.info(f"Добавлено {len(all_ids)} чанков, пропущено {skipped_count} дубликатов")
            else:
                logger.info(f"Все чанки уже существуют в базе, обновлений не требуется")

        except Exception as e:
            logger.error(f"Ошибка при добавлении документов в векторную базу: {e}")
            raise

    def search_documents(self,
                         query: str,
                         n_results: int = None,
                         filter_metadata: Dict = None) -> List[Dict[str, Any]]:
        """
        Поиск релевантных документов по запросу

        Args:
            query: Поисковый запрос
            n_results: Количество результатов
            filter_metadata: Фильтры по метаданным

        Returns:
            Список релевантных чанков
        """
        try:
            n_results = n_results or Config.RAG_N_RESULTS

            # Выполняем поиск
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results,
                where=filter_metadata
            )

            formatted_results = []

            for i in range(len(results['ids'][0])):
                result = {
                    'id': results['ids'][0][i],
                    'document': results['documents'][0][i],
                    'metadata': results['metadatas'][0][i],
                    'distance': results['distances'][0][i] if results['distances'] else None
                }
                formatted_results.append(result)

            logger.debug(f"Найдено {len(formatted_results)} результатов для запроса: {query}")
            return formatted_results

        except Exception as e:
            logger.error(f"Ошибка при поиске документов: {e}")
            return []

    def get_collection_info(self) -> Dict[str, Any]:
        """
        Получение подробной информации о коллекции

        Returns:
            Информация о коллекции
        """
        try:
            count = self.collection.count()

            # Получаем статистику
            all_docs = self.collection.get(limit=1000)
            total_chunks = len(all_docs['ids']) if all_docs else 0

            # Статистика по файлам
            files_info = {}
            if all_docs and 'metadatas' in all_docs:
                for metadata in all_docs['metadatas']:
                    filename = metadata.get('filename', 'unknown')
                    if filename not in files_info:
                        files_info[filename] = 0
                    files_info[filename] += 1

            return {
                "collection_name": self.collection_name,
                "total_chunks": count,
                "unique_files": len(files_info),
                "files_info": files_info,
                "embedding_model": self.model.embedding_model,
                "chroma_dir": str(self.chroma_dir),
                "docs_dir": str(self.docs_dir),
                "chunk_size": Config.CHUNK_SIZE,
                "collection_created": self.collection.metadata.get("created_at")
            }
        except Exception as e:
            logger.error(f"Ошибка при получении информации о коллекции: {e}")
            return {}

    def clear_collection(self):
        """
        Очистка коллекции
        """
        try:
            old_count = self.collection.count()
            self.chroma_client.delete_collection(self.collection_name)

            # Создаем новую коллекцию
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name,
                embedding_function=self.embedding_function,
                metadata={
                    "description": "Хранилище векторных представлений документов",
                    "embedding_model": self.model.embedding_model,
                    "created_at": datetime.now().isoformat(),
                    "reset_at": datetime.now().isoformat()
                }
            )

            logger.info(f"Коллекция очищена, удалено {old_count} чанков")

        except Exception as e:
            logger.error(f"Ошибка при очистке коллекции: {e}")
            raise

    def update_document(self, file_path: Path):
        """
        Обновление документа в векторной базе

        Args:
            file_path: Путь к обновленному файлу
        """
        try:
            # Удаляем старые записи для этого файла
            file_hash = self.get_file_hash(file_path)

            # Ищем и удаляем существующие записи
            existing_results = self.collection.get(
                where={"filename": file_path.name}
            )

            if existing_results['ids']:
                self.collection.delete(ids=existing_results['ids'])
                logger.info(f"Удалено {len(existing_results['ids'])} старых чанков для {file_path.name}")

            # Добавляем обновленный документ
            doc_info = self.process_single_document(file_path)
            if doc_info:
                self.add_documents_to_vector_db([doc_info])
                logger.info(f"Документ {file_path.name} обновлен")

        except Exception as e:
            logger.error(f"Ошибка при обновлении документа {file_path}: {e}")
            raise